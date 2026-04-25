from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path("memoria_madrid_refugio.docx")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    run.font.size = Pt(11)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Lato")
    r_fonts.set(qn("w:hAnsi"), "Lato")
    r_fonts.set(qn("w:eastAsia"), "Lato")
    r_pr.append(r_fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "22")
    r_pr.append(size_cs)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = document.styles
    if "BodyLato" not in styles:
        style = styles.add_style("BodyLato", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["BodyLato"]

    style.font.name = "Lato"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)

    normal = styles["Normal"]
    normal.font.name = "Lato"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style="BodyLato")
    paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    run.font.size = Pt(12 if level == 1 else 11)


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph(style="BodyLato")
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        prefix_run.bold = True
        prefix_run.font.name = "Lato"
        prefix_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
        prefix_run.font.size = Pt(11)
        rest = text[len(bold_prefix):]
        run = paragraph.add_run(rest)
        run.font.name = "Lato"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
        run.font.size = Pt(11)
        return

    run = paragraph.add_run(text)
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    run.font.size = Pt(11)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="BodyLato")
        paragraph.paragraph_format.left_indent = Cm(0.63)
        paragraph.paragraph_format.first_line_indent = Cm(-0.4)
        bullet = paragraph.add_run("• ")
        bullet.font.name = "Lato"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
        bullet.font.size = Pt(11)
        run = paragraph.add_run(item)
        run.font.name = "Lato"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
        run.font.size = Pt(11)


def build_document() -> Document:
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="BodyLato")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("MEMORIA DEL PROYECTO")
    run.bold = True
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    run.font.size = Pt(13)

    subtitle = document.add_paragraph(style="BodyLato")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Madrid Refugio: rutas de confort térmico urbano basadas en datos abiertos")
    run.bold = True
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    run.font.size = Pt(11)

    add_heading(document, "1. Resumen ejecutivo")
    add_paragraph(document, "Madrid Refugio es una herramienta de simulación climática urbana que calcula rutas peatonales de confort térmico en el momento de la consulta, combinando datos abiertos del Ayuntamiento de Madrid con un grafo de calles de OpenStreetMap. El proyecto responde a un problema público acreditado: el 64,1% de los barrios de Madrid no dispone de un refugio climático operativo en un radio de 300 metros, mientras que la cobertura actual de refugios oficiales resulta claramente insuficiente en comparación con otras grandes ciudades españolas.")
    add_paragraph(document, "Ante esa brecha de infraestructura, Madrid Refugio propone una solución operativa e inmediata: ayudar a que cada desplazamiento a pie sea más seguro durante episodios de calor extremo, maximizando la sombra disponible, la proximidad a fuentes de agua y el acceso a equipamientos climatizados que pueden funcionar como refugios sustitutos.")

    add_heading(document, "2. Problema y oportunidad pública")
    add_paragraph(document, "Las olas de calor son ya uno de los principales riesgos climáticos para la salud urbana. En Madrid, las temperaturas extremas afectan de forma desproporcionada a las personas mayores de 65 años, especialmente en barrios con menor cobertura de arbolado, mayor distancia a equipamientos de resguardo y peor acceso a recursos de alivio inmediato.")
    add_paragraph(document, "El análisis territorial realizado sobre 131 barrios pone de manifiesto tres conclusiones relevantes. En primer lugar, 84 barrios, el 64,1% del total, no cuentan con un refugio climático operativo a menos de 300 metros. En segundo lugar, Aluche concentra la mayor vulnerabilidad absoluta en volumen de población mayor expuesta, con 19.121 personas mayores de 65 años sin refugio próximo. En tercer lugar, Villaverde Alto - Casco Histórico de Villaverde obtiene la máxima prioridad relativa de intervención en el modelo territorial.")
    add_paragraph(document, "Durante el desarrollo se detectó, además, una carencia relevante del ecosistema de datos: no existe en el portal municipal un dataset específico y reutilizable de refugios climáticos oficiales. El proyecto documenta esa ausencia y la compensa de forma trazable mediante una capa operativa de equipamientos municipales climatizados.")

    add_heading(document, "3. Objetivo del proyecto")
    add_paragraph(document, "Objetivo ciudadano. Ofrecer una herramienta gratuita y accesible desde cualquier navegador que permita calcular rutas peatonales optimizadas para el confort térmico entre dos puntos de Madrid, teniendo en cuenta sombra urbana, proximidad a fuentes de agua potable y acceso a refugios sustitutos climatizados.", bold_prefix="Objetivo ciudadano. ")
    add_paragraph(document, "Objetivo de política pública. Generar una base territorial de evidencia que ayude a priorizar la ampliación de la red de refugios climáticos, la plantación de arbolado viario y otras medidas de adaptación en los barrios con mayor vulnerabilidad.", bold_prefix="Objetivo de política pública. ")

    add_heading(document, "4. Público beneficiario")
    add_bullets(document, [
        "Ciudadanía en general, especialmente personas mayores de 65 años que necesitan desplazarse a pie en episodios de calor extremo.",
        "Servicios sociales, sanitarios y de emergencia municipal, para orientar recursos y actuaciones preventivas durante alertas térmicas.",
        "Gestores públicos y técnicos municipales responsables de planificación urbana, adaptación climática y salud ambiental.",
    ])

    add_heading(document, "5. Conjuntos de datos utilizados")
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Conjunto de datos", "Fuente", "Formato", "Uso en el proyecto"]
    for idx, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], value, bold=True)

    rows = [
        ("Modelo de alturas de edificación (2024)", "Geoportal del Ayuntamiento de Madrid", "Capa geoespacial / GeoJSON procesado", "Base para el cálculo de sombra proyectada por edificios"),
        ("Arbolado viario detallado", "datos.madrid.es", "XLSX", "Sombra biológica, peso de tramo y análisis territorial"),
        ("Zonas verdes por distrito", "datos.madrid.es", "CSV", "Contexto territorial complementario"),
        ("Bibliotecas municipales", "datos.madrid.es", "GEO", "Refugios sustitutos operativos"),
        ("Centros culturales", "datos.madrid.es", "GEO", "Refugios sustitutos operativos"),
        ("Polideportivos municipales", "datos.madrid.es", "GEO", "Refugios sustitutos operativos"),
        ("Fuentes de agua potable", "datos.madrid.es", "CSV", "Puntos de alivio en ruta y análisis de proximidad"),
        ("Calidad del aire histórica 2024", "datos.madrid.es", "ZIP CSV", "Series horarias de NO2 para el índice de confort"),
        ("Estaciones de calidad del aire", "datos.madrid.es", "CSV", "Coordenadas para interpolación espacial"),
        ("Padrón municipal por edad", "datos.madrid.es", "CSV", "Población mayor de 65 años por barrio"),
        ("Límites administrativos de barrios", "Geoportal del Ayuntamiento de Madrid", "ZIP Shapefile", "Unidad territorial oficial de análisis"),
        ("Red peatonal urbana", "OpenStreetMap", "Grafo OSM", "Base de routing peatonal"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    add_paragraph(document, "Nota sobre datos no disponibles. No existe en el portal municipal un dataset específico de refugios climáticos oficiales reutilizable en formato operativo. Por ello, el proyecto trabaja con una sustitución defensiva basada en 261 equipamientos con coordenadas verificadas, manteniendo trazabilidad completa de la decisión.", bold_prefix="Nota sobre datos no disponibles. ")
    add_paragraph(document, "Fuente comparativa externa. El informe de Greenpeace de 2025 se utiliza como referencia de contexto y benchmarking sobre cobertura de refugios climáticos, no como dataset operativo del motor de cálculo.", bold_prefix="Fuente comparativa externa. ")

    add_heading(document, "6. Innovación que representa")
    add_paragraph(document, "La principal innovación del proyecto es la incorporación de un modelo de sombra urbana por hora del día aplicado al cálculo de rutas peatonales. Frente a visores estáticos o mapas agregados de calor, Madrid Refugio estima el confort térmico tramo a tramo y compara una ruta directa con una alternativa más protegida.")
    add_paragraph(document, "Esa estimación combina dos tipos de sombra. Por un lado, la sombra proyectada por edificios a partir del modelo de alturas de edificación y de la posición solar calculada con pvlib y pybdshadow. Por otro, la sombra biológica derivada del inventario municipal de arbolado viario, integrado sobre la red peatonal para modificar el peso de cada tramo.")
    add_paragraph(document, "La segunda innovación es el índice territorial multicriterio por barrio, que combina cobertura de arbolado, proximidad a fuentes, cobertura de refugios sustitutos, calidad del aire y peso de la población vulnerable. El resultado no solo ayuda a caminar con menor exposición al calor, sino que genera evidencia útil para priorizar inversión pública.")
    add_paragraph(document, "En el ejemplo operativo validado en el proyecto, una ruta de confort térmico incrementa la distancia un 4,3% respecto al trayecto más corto, a cambio de multiplicar por 5,4 la sombra acumulada.")
    add_paragraph(document, "6.1 Lógica de routing climático", bold_prefix="6.1 Lógica de routing climático")
    add_paragraph(document, "En cada consulta el backend calcula dos alternativas sobre la misma red peatonal: una ruta directa, optimizada exclusivamente por longitud, y una ruta de confort térmico, optimizada mediante una función de coste dinámica. Esa función combina cuatro factores: la longitud del tramo, la sombra biológica procedente del arbolado, la sombra proyectada por edificios para la franja horaria solicitada y un bono de proximidad a recursos de alivio.")
    add_paragraph(document, "La hora introducida por la persona usuaria se transforma en una franja discreta comprendida entre las 08:00 y las 20:00. Con ello, el motor selecciona la matriz horaria de sombra precomputada correspondiente y evita recalcular la geometría solar completa en producción. La preferencia de ruta se modela en un continuo técnico entre 0 y 1, pero se expone en tres modos comprensibles: Directa, Equilibrada y Más sombra.")
    add_paragraph(document, "6.2 Recursos de proximidad y salida operativa", bold_prefix="6.2 Recursos de proximidad y salida operativa")
    add_paragraph(document, "Las fuentes de agua y los refugios sustitutos no se incorporan como meros puntos de contexto, sino como recursos detectados espacialmente sobre la geometría de cada ruta. El sistema aplica buffers funcionales sobre el recorrido calculado y selecciona los puntos contenidos en ese corredor: 75 metros para fuentes de agua potable y 200 metros para refugios sustitutos climatizados.")
    add_paragraph(document, "La respuesta del endpoint no devuelve solo una polilínea cartográfica. Incluye las coordenadas de la ruta directa y la ruta de confort, la longitud de ambas alternativas, la sombra acumulada diferenciada entre arbolado y edificios, la reducción térmica estimada, el tiempo adicional asumido y el número de fuentes y refugios próximos. Esta salida permite auditar cada recomendación y convierte la herramienta en un sistema interpretable para ciudadanía y evaluación pública.")

    add_heading(document, "7. Tecnología utilizada")
    add_paragraph(document, "Procesamiento geoespacial. Python con GeoPandas, Shapely, PyProj, Pandas y utilidades de análisis espacial sobre ETRS89 / UTM zona 30N (EPSG:25830). Para el modelo de sombra proyectada se utilizan pvlib y pybdshadow.", bold_prefix="Procesamiento geoespacial. ")
    add_paragraph(document, "Modelo de routing. OSMnx para construir el grafo peatonal desde OpenStreetMap y NetworkX para el cálculo de rutas ponderadas por confort térmico.", bold_prefix="Modelo de routing. ")
    add_paragraph(document, "Aplicación web. Frontend en Next.js con React. API en Python con FastAPI. Visualización cartográfica con Leaflet y React-Leaflet. Despliegue del frontend en Vercel y del backend en infraestructura propia (servidor privado con Cloudflare Tunnel), con Railway como alternativa de despliegue.", bold_prefix="Aplicación web. ")
    add_paragraph(document, "Arquitectura de cálculo. El sistema combina precomputación offline de capas pesadas, como la matriz de sombra por franja horaria, con cálculo interactivo en el momento de la consulta. Este enfoque permite tiempos de respuesta operativos sin renunciar a detalle geoespacial.", bold_prefix="Arquitectura de cálculo. ")
    add_paragraph(document, "Trazabilidad. El proyecto conserva documentación técnica, scripts de generación y referencia directa a las fuentes originales utilizadas durante el procesamiento y la evaluación.", bold_prefix="Trazabilidad. ")

    add_heading(document, "8. Impacto esperado")
    add_paragraph(document, "Madrid Refugio tiene un impacto potencial directo en salud urbana, adaptación climática y planificación territorial. En el plano ciudadano, facilita desplazamientos más seguros y mejor informados en periodos de calor extremo, especialmente para población vulnerable. En el plano institucional, sistematiza evidencias sobre barrios con mayor déficit de cobertura y ofrece un soporte objetivo para decidir dónde intervenir primero.")
    add_paragraph(document, "El proyecto también contribuye a mejorar el propio ecosistema de datos abiertos, al visibilizar la ausencia de un dataset municipal de refugios climáticos y demostrar por qué ese recurso es relevante para la acción pública.")
    add_paragraph(document, "Su arquitectura y su metodología son, además, replicables en otros municipios que dispongan de datos abiertos comparables sobre red viaria, arbolado, edificios y equipamientos.")

    add_heading(document, "9. Conclusión")
    add_paragraph(document, "Madrid Refugio transforma datos abiertos en protección climática concreta. Es una herramienta operativa, desplegada y funcional, que permite calcular hoy rutas peatonales más confortables en Madrid y, al mismo tiempo, producir evidencia territorial para orientar decisiones públicas de adaptación al calor.")
    add_paragraph(document, "La combinación de reutilización de datos, modelización geoespacial y utilidad ciudadana convierte el proyecto en un ejemplo sólido de innovación aplicada al interés general.")

    paragraph = document.add_paragraph(style="BodyLato")
    prefix = paragraph.add_run("URL pública de la herramienta: ")
    prefix.bold = True
    prefix.font.name = "Lato"
    prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    prefix.font.size = Pt(11)
    add_hyperlink(paragraph, "https://madridrefugio.es/", "https://madridrefugio.es/")

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
